import streamlit as st
import json
from pathlib import Path
from datetime import datetime, timedelta, timezone
import io
import hashlib
import zlib
import zipfile

import pandas as pd
import numpy as np

# Optional libs (Fetch / Report)
try:
    import requests
except Exception:
    requests = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    canvas = None
    A4 = None

# Optional: entsoe-py (falls installiert)
try:
    from entsoe import EntsoePandasClient
except Exception:
    EntsoePandasClient = None


# ================================================================
# 0) PAGE CONFIG (muss als erstes Streamlit-Kommando kommen)
# ================================================================
st.set_page_config(page_title="Profi-Planung OS", layout="wide", page_icon="🏗️")


# ================================================================
# 1) KONSTANTEN
# ================================================================
SCOPE_OPTIONS = ["PV-System", "Speicher", "Ladeinfrastruktur"]
HAK_OPTIONS = [35, 50, 63, 80, 100, 125, 160, 250]
DACH_FORMS = ["Satteldach", "Flachdach", "Pultdach", "Walmdach"]
BAT_TYPES = ["Hochvolt", "Niedervolt"]
PV_KONZEPTE = ["Überschusseinspeisung", "Volleinspeisung"]
FZ_KLASSEN = ["PKW", "Transporter", "LKW"]
SIM_RES_OPTIONS = ["15min", "hour"]

# Heuristik: kWh/kWp/a (grob)
REGIONS = {"Nord": 940, "Mitte": 1060, "Süd": 1180}

OBJECTIVES = [
    "Eigenverbrauch/Autarkie",
    "Depotcharging/Lastmanagement",
    "Arbitrage (Stromhandel)",
    "Kombi (Eigenverbrauch + Arbitrage)",
]
HAK_MODE_OPTIONS = ["Ampere", "kW"]

CHARGER_TYPES = ["AC", "DC"]
CONSUMPTION_MODE = ["Manuell", "KI (Durchschnitt)"]

# PV-Ertrag Modell (neu)
PV_YIELD_MODE_OPTIONS = ["Heuristik (Region)", "PVGIS (Standort)"]
PVGIS_URL = "https://re.jrc.ec.europa.eu/api/"

# KI-Durchschnittswerte (grob, praxisnah; anpassbar)
CONSUMPTION_DEFAULTS = {
    "PKW": {"kwh_per_100km": 18.0, "l_per_100km": 7.0},
    "Transporter": {"kwh_per_100km": 28.0, "l_per_100km": 10.0},
    "LKW": {"kwh_per_100km": 120.0, "l_per_100km": 30.0},
}

# Plausibilitäts-Grenzen (weich, Warnung statt Blockade)
PLAUS_ROOF_W_MAX = 300.0          # m
PLAUS_ROOF_D_MAX = 300.0          # m
PLAUS_ROOF_AREA_WARN = 10_000.0   # m²
PLAUS_ROOF_AREA_HARD = 1_000_000.0
PLAUS_ROOF_MOD_WARN = 200_000     # Module


# ================================================================
# 2) DEFAULTS / MIGRATION
# ================================================================
def _hak_kw_from_ampere(a: float) -> float:
    return (float(a) * 400.0 * 1.73) / 1000.0


def _hak_ampere_from_kw(kw: float) -> int:
    a = float(kw) * 1000.0 / (400.0 * 1.73)
    return int(max(round(a), 1))


DEFAULTS = {
    "kunde": {"name": "—"},
    "scope": ["PV-System", "Speicher", "Ladeinfrastruktur"],
    "tech": {
        "hak_mode": "Ampere",
        "hak_ampere": 63,
        "hak_kw": round(_hak_kw_from_ampere(63), 1),
        "ems_required": False,
        "berechnungs_modus": "Detail",
    },
    "dach": {"form": "Satteldach", "breite": 10.0, "tiefe": 6.0, "flaeche": 60.0, "neigung": 35, "azimut": 0},
    "daecher": [
        {
            "id": "roof-1",
            "name": "Dachfläche 1",
            "form": "Satteldach",
            "breite": 10.0,
            "tiefe": 6.0,
            "flaeche": 60.0,
            "flaeche_auto": True,
            "neigung": 35,
            "azimut": 0,
            "nutzfaktor": 0.80,
            "module_area_m2": 2.0,
            "pv_kwp_manual_enabled": False,
            "pv_kwp_manual": 0.0,
            "hinweis": "",
        }
    ],
    "pv": {
        "felder": [],  # roof_id, hersteller, typ, watt_pro_modul, anzahl_module, datasheet
        "wr": [],      # modell, ac_kw, datasheet
        "konzept": "Überschusseinspeisung",
        "total_kwp": 0.0,
        "total_ac_kw": 0.0,
        # Neu: PV-Ertragsmodell
        "yield": {
            "mode": "Heuristik (Region)",
            "region": "Mitte",
            # PVGIS
            "location_query": "",
            "lat": 51.1657,    # DE Mitte grob
            "lon": 10.4515,
            "loss_pct": 14.0,  # typische Systemverluste
            "ref_year": 2020,  # Referenzjahr (PVGIS)
            "use_roof_angles": True,
            "tilt_default": 30.0,
            "azimuth_default": 0.0,  # 0=south, 90=west, -90=east (PVGIS Konvention)
            "mountingplace": "free",  # 'free' oder 'building'
            "raddatabase": "",       # optional leer
        },
    },
    "speicher": {
        "hersteller": "",
        "kap": 0.0,
        "p": 0.0,
        "spannung": "Hochvolt",
        "datasheet": "",
        "objective": "Eigenverbrauch/Autarkie",
        "min_soc_pct": 10.0,
        "cycle_life": 6000,
        "calendar_life_years": 15,
        "warranty_years": 10,
        "eta_roundtrip": 0.92,
        "arbitrage": {
            "enabled": False,
            "mode": "Manuell",
            "low_price": 0.18,
            "high_price": 0.32,
            "spread": 0.14,
            "cycles_per_year": 180,
            "roundtrip_eff": 0.90,
            "dod": 0.90,
            "degradation_eur_per_kwh_throughput": 0.03,
            "source": "SMARD (DE/LU)",
            "smard_series_id": 8004169,
            "smard_region": "DE-LU",
            "lookback_days": 30,
            "lookahead_hours": 0,
            "low_quantile": 0.20,
            "high_quantile": 0.80,
            "entsoe_token": "",
            "entsoe_bidding_zone": "DE_LU",
            "price_csv_file": "",
            "price_csv_unit": "EUR/kWh",
            "continuous_soc": True,
            "soc_grid_points": 201,
            "terminal_mode": "End-SOC = SOC_min",
            "terminal_value_mode": "Auto (letzter Preis)",
            "terminal_value_eur_per_kwh": 0.25,
            "prices_cached": False,
            "prices_last_fetch": "",
            "schedule_file": "",
            "annual_profit_est": 0.0,
            "annual_throughput_kwh": 0.0,
            "annual_cycles_est": 0.0,
        },
    },
    "mobilität": {
        "ladepunkte": [
            # {"name":"", "typ":"AC/DC", "leistung_kw_pro_punkt":11.0, "anzahl_punkte":1, "datasheet":""}
        ],
        "fuhrpark": [
            # {"klasse":"PKW", "anzahl":1, "km_pro_jahr":20000,
            #  "consumption_mode":"Manuell/KI", "kwh_pro_100km":18.0, "l_pro_100km":7.0, "datasheet":""}
        ],
        "total_ev_kwh": 0.0,
    },
    "wirtschaft": {
        "strompreis": 0.35,
        "dieselpreis": 1.70,
        "einspeise_v": 0.08,
        "lastgang_jahr": 5000,
        "lastgang_file": "",
        "lastgang_datetime_col": "",
        "lastgang_value_col": "",
        "lastgang_sep": ";",
        "lastgang_decimal": ",",
        "lastgang_resolution": "15min",
        "lastgang_is_power_kw": True,
        "capex_pv_eur_per_kwp": 1100.0,
        "capex_bat_eur_per_kwh": 450.0,
        "capex_bat_eur_per_kw": 200.0,
        "opex_pv_pct": 0.015,
        "opex_bat_pct": 0.01,
    },
}


def deep_merge(target: dict, defaults: dict) -> None:
    for k, v in defaults.items():
        if k not in target:
            target[k] = v
        else:
            if isinstance(v, dict) and isinstance(target.get(k), dict):
                deep_merge(target[k], v)


def sanitize(obj):
    if isinstance(obj, dict):
        return {k: sanitize(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [sanitize(v) for v in obj]
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating,)):
        x = float(obj)
        if np.isnan(x) or np.isinf(x):
            return 0.0
        return x
    if isinstance(obj, float):
        if np.isnan(obj) or np.isinf(obj):
            return 0.0
        return obj
    return obj


def dict_hash(obj: dict) -> str:
    s = json.dumps(sanitize(obj), sort_keys=True, ensure_ascii=False).encode("utf-8")
    return hashlib.md5(s).hexdigest()


def stable_seed(*parts) -> int:
    payload = "|".join(map(str, parts)).encode("utf-8")
    return zlib.crc32(payload) & 0xFFFFFFFF


def safe_index(options, value, fallback=0) -> int:
    try:
        return options.index(value)
    except ValueError:
        return fallback


def hak_p_max_kw_from_state(d: dict) -> float:
    mode = str(d.get("tech", {}).get("hak_mode", "Ampere"))
    if mode == "kW":
        return float(d.get("tech", {}).get("hak_kw", _hak_kw_from_ampere(d["tech"].get("hak_ampere", 63))) or 0.0)
    return _hak_kw_from_ampere(float(d.get("tech", {}).get("hak_ampere", 63) or 63))


def ensure_roof_ids(d: dict):
    for i, r in enumerate(d.get("daecher", [])):
        if not r.get("id"):
            r["id"] = f"roof-{i+1}"


def _migrate_mobility(d: dict):
    mob = d.setdefault("mobilität", {})
    deep_merge(mob, DEFAULTS["mobilität"])

    lp_new = []
    for lp in mob.get("ladepunkte", []):
        if not isinstance(lp, dict):
            continue
        if "leistung_kw_pro_punkt" in lp and "anzahl_punkte" in lp:
            lp.setdefault("typ", "DC" if float(lp.get("leistung_kw_pro_punkt", 0) or 0) >= 50 else "AC")
            lp.setdefault("name", lp.get("name", "Ladepunkt"))
            lp.setdefault("datasheet", lp.get("datasheet", ""))
            lp_new.append(lp)
        else:
            name = lp.get("name", lp.get("n", "Ladepunkt"))
            p = lp.get("leistung_kw", lp.get("p", 11))
            try:
                p = float(p)
            except Exception:
                p = 11.0
            typ = lp.get("typ")
            if not typ:
                typ = "DC" if p >= 50 else "AC"
            lp_new.append({
                "name": name,
                "typ": typ,
                "leistung_kw_pro_punkt": p,
                "anzahl_punkte": int(lp.get("anzahl_punkte", 1) or 1),
                "datasheet": lp.get("datasheet", ""),
            })
    mob["ladepunkte"] = lp_new

    fz_new = []
    for fz in mob.get("fuhrpark", []):
        if not isinstance(fz, dict):
            continue
        if "km_pro_jahr" in fz and "kwh_pro_100km" in fz and "l_pro_100km" in fz:
            fz.setdefault("consumption_mode", "Manuell")
            fz.setdefault("datasheet", fz.get("datasheet", ""))
            fz_new.append(fz)
        else:
            klasse = fz.get("klasse", fz.get("typ", "PKW"))
            if klasse not in FZ_KLASSEN:
                klasse = "PKW"
            fz_new.append({
                "klasse": klasse,
                "anzahl": int(fz.get("anz", fz.get("anzahl", 1)) or 1),
                "km_pro_jahr": int(fz.get("km", fz.get("km_pro_jahr", 20000)) or 0),
                "consumption_mode": "Manuell",
                "kwh_pro_100km": float(fz.get("e_cons", fz.get("kwh_pro_100km", CONSUMPTION_DEFAULTS[klasse]["kwh_per_100km"])) or 0.0),
                "l_pro_100km": float(fz.get("d_cons", fz.get("l_pro_100km", CONSUMPTION_DEFAULTS[klasse]["l_per_100km"])) or 0.0),
                "datasheet": fz.get("datasheet", ""),
            })
    mob["fuhrpark"] = fz_new


def migrate_legacy(d: dict):
    d.setdefault("tech", {})
    deep_merge(d["tech"], DEFAULTS["tech"])
    if not d["tech"].get("hak_kw"):
        d["tech"]["hak_kw"] = round(_hak_kw_from_ampere(float(d["tech"].get("hak_ampere", 63))), 1)

    if "daecher" not in d or not isinstance(d["daecher"], list) or len(d["daecher"]) == 0:
        old = d.get("dach", {})
        d["daecher"] = [{
            "id": "roof-1",
            "name": "Dachfläche 1",
            "form": old.get("form", "Satteldach"),
            "breite": float(old.get("breite", 10.0)),
            "tiefe": float(old.get("tiefe", 6.0)),
            "flaeche": float(old.get("flaeche", 60.0)),
            "flaeche_auto": True,
            "neigung": int(old.get("neigung", 35)),
            "azimut": int(old.get("azimut", 0)),
            "nutzfaktor": 0.80,
            "module_area_m2": 2.0,
            "pv_kwp_manual_enabled": False,
            "pv_kwp_manual": 0.0,
            "hinweis": "",
        }]
    ensure_roof_ids(d)

    d.setdefault("pv", {})
    deep_merge(d["pv"], DEFAULTS["pv"])
    d["pv"].setdefault("yield", {})
    deep_merge(d["pv"]["yield"], DEFAULTS["pv"]["yield"])

    wr_new = []
    for w in d.get("pv", {}).get("wr", []):
        if isinstance(w, dict) and ("ac_kw" in w or "modell" in w):
            w.setdefault("modell", w.get("h", ""))
            w.setdefault("ac_kw", float(w.get("p", 0.0) or 0.0))
            w.setdefault("datasheet", w.get("datasheet", ""))
            wr_new.append(w)
        elif isinstance(w, dict):
            wr_new.append({
                "modell": w.get("h", ""),
                "ac_kw": float(w.get("p", 0.0) or 0.0),
                "datasheet": w.get("datasheet", ""),
            })
    d["pv"]["wr"] = wr_new

    fields_new = []
    for f in d.get("pv", {}).get("felder", []):
        if not isinstance(f, dict):
            continue
        if "watt_pro_modul" in f:
            if not f.get("roof_id") and d.get("daecher"):
                f["roof_id"] = d["daecher"][0]["id"]
            f.setdefault("hersteller", f.get("h", ""))
            f.setdefault("typ", f.get("t", ""))
            f.setdefault("anzahl_module", int(f.get("n", 0) or 0))
            f.setdefault("datasheet", f.get("datasheet", ""))
            fields_new.append(f)
        else:
            rid = d["daecher"][0]["id"]
            fields_new.append({
                "roof_id": f.get("roof_id", rid),
                "hersteller": f.get("h", ""),
                "typ": f.get("t", ""),
                "watt_pro_modul": float(f.get("w", 0) or 0),
                "anzahl_module": int(f.get("n", 0) or 0),
                "datasheet": f.get("datasheet", ""),
            })
    d["pv"]["felder"] = fields_new

    d.setdefault("speicher", {})
    deep_merge(d["speicher"], DEFAULTS["speicher"])
    d.setdefault("wirtschaft", {})
    deep_merge(d["wirtschaft"], DEFAULTS["wirtschaft"])

    _migrate_mobility(d)


def load_project(slug: str) -> tuple[dict, Path]:
    p_path = Path("projects") / slug
    state_file = p_path / "state.json"
    p_path.mkdir(parents=True, exist_ok=True)

    if not state_file.exists():
        st.error("⚠️ Projektdatei nicht gefunden (state.json). Bitte Projekt über Dashboard anlegen/öffnen.")
        st.stop()

    with state_file.open("r", encoding="utf-8") as f:
        data = json.load(f)

    deep_merge(data, DEFAULTS)
    migrate_legacy(data)
    return data, state_file


def save_project(state_file: Path, data: dict):
    state_file.parent.mkdir(parents=True, exist_ok=True)
    with state_file.open("w", encoding="utf-8") as f:
        json.dump(sanitize(data), f, indent=4, ensure_ascii=False)


def save_if_changed(state_file: Path, data: dict):
    h = dict_hash(data)
    if st.session_state.get("_last_saved_hash") != h:
        save_project(state_file, data)
        st.session_state["_last_saved_hash"] = h


# ================================================================
# 3) PATHS / FILE HELPERS
# ================================================================
def project_paths(slug: str):
    base = Path("projects") / slug
    docs = base / "documents"
    docs.mkdir(parents=True, exist_ok=True)
    ds = docs / "datasheets"
    ds.mkdir(parents=True, exist_ok=True)
    inputs = docs / "inputs"
    inputs.mkdir(parents=True, exist_ok=True)
    reports = docs / "reports"
    reports.mkdir(parents=True, exist_ok=True)
    arb = docs / "arbitrage"
    arb.mkdir(parents=True, exist_ok=True)
    return base, docs, ds, inputs, reports, arb


def save_uploaded_file(folder: Path, uploaded_file, prefix: str = "") -> str:
    folder.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = uploaded_file.name.replace("/", "_").replace("\\", "_")
    filename = f"{prefix}{ts}_{safe_name}"
    path = folder / filename
    path.write_bytes(uploaded_file.getbuffer())
    return str(path)


def export_project_zip_bytes(slug: str) -> bytes:
    """Zipped das komplette Projektverzeichnis (inkl. documents) in-memory."""
    proj_dir = Path("projects") / slug
    if not proj_dir.exists():
        return b""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in proj_dir.rglob("*"):
            if p.is_file():
                # arcname so, dass beim Entpacken wieder projects/<slug>/... entsteht
                arcname = str(Path("projects") / slug / p.relative_to(proj_dir)).replace("\\", "/")
                z.write(p, arcname=arcname)
    return buf.getvalue()


def import_project_zip_bytes(zip_bytes: bytes) -> str | None:
    """Importiert ZIP in ./projects. Gibt den gefundenen Slug zurück (best-effort)."""
    projects_root = Path("projects")
    projects_root.mkdir(parents=True, exist_ok=True)

    slug_guess = None
    with zipfile.ZipFile(io.BytesIO(zip_bytes), "r") as z:
        names = z.namelist()
        # Versuche slug aus erstem projects/<slug>/... Pfad zu erkennen
        for n in names:
            parts = n.replace("\\", "/").split("/")
            if len(parts) >= 2 and parts[0] == "projects" and parts[1]:
                slug_guess = parts[1]
                break
        z.extractall(".")  # entpackt projects/... korrekt, wenn arcname so gepackt war

    # Fallback: falls slug nicht im ZIP stand, suche neueste Ordner
    if slug_guess and (projects_root / slug_guess).exists():
        return slug_guess

    dirs = [p for p in projects_root.iterdir() if p.is_dir()]
    if not dirs:
        return None
    dirs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return dirs[0].name


# ================================================================
# 4) ROOF / PV CALCS
# ================================================================
def roof_area(roof: dict) -> float:
    b = float(roof.get("breite", 0.0))
    t = float(roof.get("tiefe", 0.0))
    return max(b * t, 0.0)


def roof_module_capacity(roof: dict) -> int:
    a = float(roof.get("flaeche", 0.0))
    nf = float(roof.get("nutzfaktor", 0.8))
    ma = float(roof.get("module_area_m2", 2.0))
    return int((a * nf) / max(ma, 0.1))


def pv_kwp_from_fields(fields: list[dict], roof_id: str = "") -> float:
    total = 0.0
    for f in fields:
        if roof_id and f.get("roof_id") != roof_id:
            continue
        w = float(f.get("watt_pro_modul", 0) or 0)
        n = float(f.get("anzahl_module", 0) or 0)
        total += (w * n) / 1000.0
    return float(total)


def modules_from_fields(fields: list[dict], roof_id: str = "") -> int:
    total = 0
    for f in fields:
        if roof_id and f.get("roof_id") != roof_id:
            continue
        total += int(f.get("anzahl_module", 0) or 0)
    return total


def pv_kwp_total(d: dict) -> float:
    fields = d["pv"]["felder"]
    total = 0.0
    for r in d["daecher"]:
        if bool(r.get("pv_kwp_manual_enabled")):
            total += float(r.get("pv_kwp_manual", 0.0) or 0.0)
        else:
            total += pv_kwp_from_fields(fields, r["id"])
    return float(total)


def wr_ac_total(d: dict) -> float:
    total = 0.0
    for w in d["pv"]["wr"]:
        total += float(w.get("ac_kw", 0.0) or 0.0)
    return float(total)


# ================================================================
# 4b) PVGIS (Standort) – Ertragssimulation (neu)
# ================================================================
@st.cache_data(show_spinner=False)
def geocode_nominatim(query: str) -> dict:
    if requests is None:
        raise RuntimeError("requests ist nicht verfügbar (pip install requests).")
    q = (query or "").strip()
    if not q:
        raise ValueError("Leere Standortabfrage.")
    url = "https://nominatim.openstreetmap.org/search"
    params = {"format": "json", "q": q, "limit": 1, "addressdetails": 1}
    headers = {"User-Agent": "EnergyOS/2.6 (Streamlit)"}
    r = requests.get(url, params=params, headers=headers, timeout=20)
    r.raise_for_status()
    data = r.json()
    if not data:
        raise ValueError("Keine Koordinaten gefunden.")
    hit = data[0]
    return {
        "lat": float(hit["lat"]),
        "lon": float(hit["lon"]),
        "display_name": hit.get("display_name", q),
    }


@st.cache_data(show_spinner=False)
def pvgis_power_per_kwp_hourly(
    lat: float,
    lon: float,
    year: int,
    tilt: float,
    azimuth: float,
    loss_pct: float,
    mountingplace: str = "free",
    raddatabase: str | None = None,
    timeout: int = 30,
) -> pd.Series:
    """
    PVGIS seriescalc: liefert stündliche PV-Leistung P (W) für peakpower=1 kW.
    Index wird auf Europe/Berlin konvertiert und tz-naiv zurückgegeben.
    PVGIS Azimut-Konvention: 0=south, 90=west, -90=east.
    """
    if requests is None:
        raise RuntimeError("requests ist nicht verfügbar (pip install requests).")

    params = {
        "lat": float(lat),
        "lon": float(lon),
        "outputformat": "json",
        "angle": float(tilt),
        "aspect": float(azimuth),
        "pvcalculation": 1,
        "peakpower": 1.0,         # 1 kWp
        "loss": float(loss_pct),
        "mountingplace": mountingplace,
        "pvtechchoice": "crystSi",
        "trackingtype": 0,
        "components": 0,
        "usehorizon": 1,
        "startyear": int(year),
        "endyear": int(year),
    }
    if raddatabase:
        params["raddatabase"] = str(raddatabase)

    headers = {"User-Agent": "EnergyOS/2.6 (Streamlit)"}
    res = requests.get(PVGIS_URL + "seriescalc", params=params, headers=headers, timeout=timeout)
    if not res.ok:
        try:
            msg = res.json().get("message")
        except Exception:
            res.raise_for_status()
        raise RuntimeError(msg or f"PVGIS Fehler: HTTP {res.status_code}")

    src = res.json()
    hourly = src.get("outputs", {}).get("hourly", [])
    if not hourly:
        raise RuntimeError("PVGIS: keine hourly-Daten erhalten.")
    df = pd.DataFrame(hourly)
    if "time" not in df.columns or "P" not in df.columns:
        raise RuntimeError("PVGIS: unerwartetes Datenformat (time/P fehlt).")

    idx = pd.to_datetime(df["time"], format="%Y%m%d:%H%M", utc=True)
    idx = idx.tz_convert("Europe/Berlin").tz_localize(None)
    power_kw = pd.to_numeric(df["P"], errors="coerce").fillna(0.0).astype(float) / 1000.0  # kW

    s = pd.Series(index=idx, data=power_kw.values).sort_index()
    s = s[~s.index.duplicated(keep="last")]
    return s


def _map_index_to_year(idx: pd.DatetimeIndex, year: int) -> pd.DatetimeIndex:
    out = []
    for ts in idx:
        try:
            out.append(ts.replace(year=year))
        except ValueError:
            # 29.02 -> 28.02
            out.append(ts.replace(year=year, day=28))
    return pd.DatetimeIndex(out)


def build_pvgis_pv_energy_profile(index: pd.DatetimeIndex, d: dict) -> pd.Series:
    """
    Baut PV kWh/Step für das Projekt auf Basis PVGIS (Standort) und den Dachparametern.
    Summiert pro Dach (kWp je Dach, Tilt/Azimut je Dach optional).
    """
    y = d.get("pv", {}).get("yield", {}) or {}
    lat = float(y.get("lat", 0.0) or 0.0)
    lon = float(y.get("lon", 0.0) or 0.0)
    if abs(lat) < 1e-6 and abs(lon) < 1e-6:
        raise ValueError("PVGIS: lat/lon fehlen.")

    ref_year = int(y.get("ref_year", 2020) or 2020)
    loss = float(y.get("loss_pct", 14.0) or 14.0)
    use_roof_angles = bool(y.get("use_roof_angles", True))
    tilt_default = float(y.get("tilt_default", 30.0) or 30.0)
    az_default = float(y.get("azimuth_default", 0.0) or 0.0)
    mountingplace = str(y.get("mountingplace", "free") or "free")
    raddb = str(y.get("raddatabase", "") or "").strip() or None

    # PV-kWp je Dach
    roof_kwps = []
    for r in d.get("daecher", []):
        if bool(r.get("pv_kwp_manual_enabled")):
            kwp = float(r.get("pv_kwp_manual", 0.0) or 0.0)
        else:
            kwp = pv_kwp_from_fields(d.get("pv", {}).get("felder", []), r.get("id", ""))
        if kwp <= 0:
            continue
        tilt = float(r.get("neigung", tilt_default) or tilt_default) if use_roof_angles else tilt_default
        az = float(r.get("azimut", az_default) or az_default) if use_roof_angles else az_default
        roof_kwps.append((kwp, tilt, az))

    if not roof_kwps:
        return pd.Series(index=index, data=np.zeros(len(index)))

    # Summiere Leistung (kW) stündlich im Referenzjahr
    power_total = None
    for kwp, tilt, az in roof_kwps:
        s_per_kwp = pvgis_power_per_kwp_hourly(
            lat=lat, lon=lon, year=ref_year, tilt=tilt, azimuth=az, loss_pct=loss,
            mountingplace=mountingplace, raddatabase=raddb
        )
        s = s_per_kwp * float(kwp)
        power_total = s if power_total is None else power_total.add(s, fill_value=0.0)

    if power_total is None or power_total.empty:
        return pd.Series(index=index, data=np.zeros(len(index)))

    # Mappe Zielindex auf Referenzjahr (typischer Jahresgang)
    idx_ref = _map_index_to_year(pd.DatetimeIndex(index).tz_localize(None), ref_year)

    # Interpoliere Leistung auf Zielauflösung im Ref-Index
    union = power_total.index.union(idx_ref)
    p_interp = power_total.reindex(union).sort_index().interpolate(method="time").reindex(idx_ref).fillna(0.0)

    # kWh/Step
    if len(index) > 1:
        dt_h = (index[1] - index[0]).total_seconds() / 3600.0
    else:
        dt_h = 1.0
    energy = p_interp.values * float(dt_h)
    return pd.Series(index=index, data=energy)


# ================================================================
# 5) LASTGANG READ (OHNE UI) – nutzt Daten aus state.json
# ================================================================
@st.cache_data(show_spinner=False)
def _read_csv_bytes(content: bytes, sep: str, decimal: str) -> pd.DataFrame:
    return pd.read_csv(io.BytesIO(content), sep=sep, decimal=decimal)


def _ensure_datetime_index(df: pd.DataFrame, dt_col: str) -> pd.DataFrame:
    out = df.copy()
    out[dt_col] = pd.to_datetime(out[dt_col], errors="coerce", utc=False)
    out = out.dropna(subset=[dt_col]).sort_values(dt_col).set_index(dt_col)
    return out


def _resample_series(s: pd.Series, resolution: str) -> pd.Series:
    if resolution == "15min":
        return s.resample("15T").mean().interpolate(limit_direction="both")
    return s.resample("H").mean().interpolate(limit_direction="both")


def load_lastgang_from_state(d: dict) -> pd.Series | None:
    w = d.get("wirtschaft", {})
    f = str(w.get("lastgang_file", "") or "").strip()
    dt_col = str(w.get("lastgang_datetime_col", "") or "").strip()
    val_col = str(w.get("lastgang_value_col", "") or "").strip()
    if not f or not dt_col or not val_col:
        return None
    p = Path(f)
    if not p.exists():
        return None

    sep = str(w.get("lastgang_sep", ";") or ";")
    dec = str(w.get("lastgang_decimal", ",") or ",")
    resolution = str(w.get("lastgang_resolution", "15min") or "15min")
    is_kw = bool(w.get("lastgang_is_power_kw", True))

    try:
        content = p.read_bytes()
        df = _read_csv_bytes(content, sep=sep, decimal=dec)
        if dt_col not in df.columns or val_col not in df.columns:
            return None
        df2 = _ensure_datetime_index(df, dt_col)
        s = pd.to_numeric(df2[val_col], errors="coerce").fillna(0.0)
        s = _resample_series(s, resolution)

        # kW -> kWh/step
        if len(s.index) > 1:
            dt_h = (s.index[1] - s.index[0]).total_seconds() / 3600.0
        else:
            dt_h = 1.0
        if is_kw:
            s = s * dt_h

        return s
    except Exception:
        return None


# ================================================================
# 6) SIMULATION (Lastgang real, PV synthetisch)
# ================================================================
def build_synth_pv_profile(index: pd.DatetimeIndex, pv_kwp: float, region: str, seed: int) -> pd.Series:
    rng = np.random.default_rng(seed)
    if pv_kwp <= 0 or len(index) == 0:
        return pd.Series(index=index, data=np.zeros(len(index)))

    spec = REGIONS.get(region, 1000)
    pv_year = pv_kwp * spec

    n = len(index)
    doy = index.dayofyear.values
    year_frac = doy / 365.0
    seasonal = (np.sin(np.pi * (year_frac - 0.08)) ** 2 + 0.08)

    hour = index.hour.values + index.minute.values / 60.0
    daily = np.maximum(np.sin(np.pi * (hour - 6) / 12), 0) ** 1.6
    clouds = np.clip(rng.normal(1.0, 0.18, n), 0.3, 1.7)

    raw = seasonal * daily * clouds
    raw_sum = raw.sum()
    pv = (raw / raw_sum) * pv_year if raw_sum > 0 else np.zeros(n)
    return pd.Series(index=index, data=pv)


def build_ev_profile(index: pd.DatetimeIndex, ev_kwh_year: float, seed: int) -> pd.Series:
    rng = np.random.default_rng(seed + 1337)
    if ev_kwh_year <= 0 or len(index) == 0:
        return pd.Series(index=index, data=np.zeros(len(index)))

    n = len(index)
    hour = index.hour.values + index.minute.values / 60.0
    evening = ((hour >= 17) & (hour <= 23)).astype(float)
    midday = ((hour >= 11) & (hour <= 15)).astype(float) * 0.4
    shape = evening + midday
    noise = np.clip(rng.normal(1.0, 0.35, n), 0.2, 2.0)
    raw = shape * noise
    s = raw.sum()
    ev = (raw / s) * ev_kwh_year if s > 0 else np.zeros(n)
    return pd.Series(index=index, data=ev)


def build_synth_load(index: pd.DatetimeIndex, annual_kwh: float, seed: int) -> pd.Series:
    rng = np.random.default_rng(seed + 777)
    n = len(index)
    if n == 0:
        return pd.Series(index=index, data=np.zeros(0))
    annual_kwh = max(float(annual_kwh), 0.0)

    doy = index.dayofyear.values
    year_frac = doy / 365.0
    seasonal = 1.0 + 0.15 * np.cos(2 * np.pi * year_frac)

    hour = index.hour.values + index.minute.values / 60.0
    daily = 0.9 + 0.35 * np.cos(2 * np.pi * (hour - 18) / 24)
    noise = np.clip(rng.normal(1.0, 0.08, n), 0.7, 1.3)

    raw = seasonal * daily * noise
    raw_sum = raw.sum()
    kwh_step = (raw / raw_sum) * annual_kwh if raw_sum > 0 else np.zeros(n)
    return pd.Series(index=index, data=kwh_step)


def simulate_timeseries(
    pv_kwh: pd.Series,
    load_kwh: pd.Series,
    bat_kwh: float,
    bat_kw: float,
    roundtrip_eff: float,
    min_soc_pct: float,
) -> pd.DataFrame:
    idx = load_kwh.index
    pv_kwh = pv_kwh.reindex(idx).fillna(0.0)
    load_kwh = load_kwh.reindex(idx).fillna(0.0)

    dt_h = (idx[1] - idx[0]).total_seconds() / 3600.0 if len(idx) > 1 else 1.0
    pmax_kwh_step = max(bat_kw, 0.0) * dt_h

    cap = max(bat_kwh, 0.0)
    reserve = cap * (float(min_soc_pct) / 100.0) if cap > 0 else 0.0
    soc = reserve

    eff = float(np.clip(roundtrip_eff, 0.5, 0.99))
    eff_c = np.sqrt(eff)
    eff_d = np.sqrt(eff)

    direct, charge, discharge, imp, exp, soc_s, served = ([] for _ in range(7))

    for p, l in zip(pv_kwh.values, load_kwh.values):
        d = min(p, l)
        surplus = p - d
        deficit = l - d

        ch = 0.0
        if cap > 0 and pmax_kwh_step > 0 and surplus > 0:
            ch_possible = min(surplus, pmax_kwh_step)
            ch_cap = max(cap - soc, 0.0)
            ch = min(ch_possible, ch_cap)
            soc += ch * eff_c
            surplus -= ch

        dis_delivered = 0.0
        if cap > 0 and pmax_kwh_step > 0 and deficit > 0 and soc > reserve:
            dis_possible = min(deficit, pmax_kwh_step)
            need_from_soc = dis_possible / eff_d
            available = max(soc - reserve, 0.0)
            take = min(need_from_soc, available)
            dis_delivered = take * eff_d
            soc -= take
            deficit -= dis_delivered

        direct.append(d)
        charge.append(ch)
        discharge.append(dis_delivered)
        exp.append(max(surplus, 0.0))
        imp.append(max(deficit, 0.0))
        soc_s.append(soc)
        served.append(d + dis_delivered)

    return pd.DataFrame(
        index=idx,
        data={
            "PV": pv_kwh.values,
            "Last": load_kwh.values,
            "Direct": direct,
            "Charge": charge,
            "Discharge": discharge,
            "Import": imp,
            "Export": exp,
            "SoC": soc_s,
            "Served": served,
        },
    )


# ================================================================
# 7) ARBITRAGE: FETCH + OPTIMIERUNG
# ================================================================
def _to_hourly_prices(prices: pd.Series) -> pd.Series:
    s = prices.copy().dropna()
    if s.empty:
        return s
    if not isinstance(s.index, pd.DatetimeIndex):
        raise ValueError("Preisserie hat keinen DatetimeIndex.")
    s = s.sort_index().resample("H").mean().interpolate(limit_direction="both")
    return s


@st.cache_data(show_spinner=False)
def fetch_smard_day_ahead_prices(series_id: int, region: str, resolution: str = "hour", lookback_days: int = 30) -> pd.Series:
    if requests is None:
        raise RuntimeError("requests ist nicht verfügbar (pip install requests).")

    base = "https://www.smard.de/app/chart_data"
    idx_url = f"{base}/{series_id}/{region}/index_{resolution}.json"
    r = requests.get(idx_url, timeout=30)
    r.raise_for_status()
    idx_json = r.json()

    timestamps = idx_json.get("timestamps") if isinstance(idx_json, dict) else idx_json
    if not timestamps:
        raise RuntimeError("Keine Timestamps von SMARD erhalten.")

    now_ms = int(datetime.now(tz=timezone.utc).timestamp() * 1000)
    from_ms = now_ms - int(lookback_days * 24 * 3600 * 1000)

    ts_candidates = [t for t in timestamps if isinstance(t, int)]
    ts_candidates.sort()
    start_ts = ts_candidates[0]
    for t in ts_candidates:
        if t <= from_ms:
            start_ts = t
        else:
            break

    data_url = f"{base}/{series_id}/{region}/{series_id}_{region}_{resolution}_{start_ts}.json"
    r2 = requests.get(data_url, timeout=30)
    r2.raise_for_status()
    data_json = r2.json()

    points = []
    if isinstance(data_json, dict):
        if "series" in data_json and isinstance(data_json["series"], list) and data_json["series"]:
            s0 = data_json["series"][0]
            if isinstance(s0, dict) and "data" in s0:
                points = s0["data"]
        elif "data" in data_json:
            points = data_json["data"]
    if not points and isinstance(data_json, list):
        points = data_json

    if not points:
        raise RuntimeError("SMARD Datenformat nicht erkannt (keine Punkte).")

    ts, vals = [], []
    for p in points:
        try:
            t, v = p[0], p[1]
            if t >= from_ms:
                ts.append(pd.to_datetime(int(t), unit="ms", utc=True).tz_convert("Europe/Berlin"))
                vals.append(float(v) / 1000.0)  # EUR/MWh -> EUR/kWh
        except Exception:
            pass

    ser = pd.Series(index=pd.DatetimeIndex(ts), data=vals).sort_index()
    ser = ser[~ser.index.duplicated(keep="last")]
    return _to_hourly_prices(ser)


@st.cache_data(show_spinner=False)
def fetch_entsoe_day_ahead_prices(token: str, bidding_zone: str, start: datetime, end: datetime) -> pd.Series:
    if EntsoePandasClient is None:
        raise RuntimeError("entsoe-py ist nicht installiert (pip install entsoe-py).")
    if not token:
        raise RuntimeError("ENTSO-E Token fehlt.")

    client = EntsoePandasClient(api_key=token)
    try:
        s = client.query_day_ahead_prices(bidding_zone, start=start, end=end)
    except Exception as e:
        raise RuntimeError(f"ENTSO-E Abfrage fehlgeschlagen: {e}")

    if hasattr(s.index, "tz") and s.index.tz is not None:
        s = s.tz_convert("Europe/Berlin")

    s = (s.astype(float) / 1000.0)  # EUR/MWh -> EUR/kWh
    return _to_hourly_prices(s.sort_index())


@st.cache_data(show_spinner=False)
def parse_price_csv(content: bytes, sep: str, decimal: str, dt_col: str, price_col: str) -> pd.Series:
    df = pd.read_csv(io.BytesIO(content), sep=sep, decimal=decimal)
    df[dt_col] = pd.to_datetime(df[dt_col], errors="coerce")
    df = df.dropna(subset=[dt_col]).sort_values(dt_col).set_index(dt_col)
    s = pd.to_numeric(df[price_col], errors="coerce").dropna()
    return _to_hourly_prices(s)


def arbitrage_fill_info_from_prices(arbd: dict, price_series: pd.Series):
    low_q = float(arbd.get("low_quantile", 0.2))
    high_q = float(arbd.get("high_quantile", 0.8))
    low = float(price_series.quantile(low_q))
    high = float(price_series.quantile(high_q))
    spread = max(high - low, 0.0)
    arbd["low_price"] = round(low, 4)
    arbd["high_price"] = round(high, 4)
    arbd["spread"] = round(spread, 4)
    arbd["prices_cached"] = True
    arbd["prices_last_fetch"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def arbitrage_profit_manual_rough(arbd: dict, bat_kwh: float) -> float:
    spread = float(arbd.get("spread", 0.0) or 0.0)
    cycles = int(arbd.get("cycles_per_year", 0) or 0)
    eff = float(arbd.get("roundtrip_eff", 0.9) or 0.9)
    dod = float(arbd.get("dod", 0.9) or 0.9)
    deg = float(arbd.get("degradation_eur_per_kwh_throughput", 0.0) or 0.0)

    usable = max(bat_kwh * dod, 0.0)
    discharged = usable * cycles
    throughput = usable * cycles * 2.0
    profit = discharged * spread * eff - throughput * deg
    return float(profit)

# (DP Optimizer bleibt unverändert – aus Platzgründen nicht kommentiert)
def _dp_optimize_horizon(
    prices_eur_per_kwh: np.ndarray,
    cap_kwh: float,
    p_kw: float,
    dod: float,
    roundtrip_eff: float,
    deg_cost_per_kwh_throughput: float,
    soc_grid_points: int,
    terminal_mode: str,
    terminal_value_eur_per_kwh: float,
) -> tuple[float, np.ndarray, np.ndarray, np.ndarray]:
    T = len(prices_eur_per_kwh)
    if T == 0:
        return 0.0, np.array([]), np.array([]), np.array([])
    if cap_kwh <= 0 or p_kw <= 0:
        soc = np.zeros(T + 1, dtype=float)
        return 0.0, soc, np.zeros(T), np.zeros(T)

    dod = float(np.clip(dod, 0.10, 1.0))
    eff = float(np.clip(roundtrip_eff, 0.50, 0.99))
    eff_c = np.sqrt(eff)
    eff_d = np.sqrt(eff)

    soc_max = float(cap_kwh)
    soc_min = float(cap_kwh * (1.0 - dod))

    p_limit_grid_kwh = float(p_kw)
    max_soc_increase = p_limit_grid_kwh * eff_c
    max_soc_decrease = (p_limit_grid_kwh / eff_d) if eff_d > 0 else 0.0

    N = max(int(soc_grid_points), 51)
    levels = np.linspace(soc_min, soc_max, N)

    neg_inf = -1e30

    if terminal_mode == "End-SOC = SOC_min":
        v_next = np.full(N, neg_inf, dtype=float)
        v_next[0] = 0.0
    elif terminal_mode == "End-SOC frei (mit Restwert)":
        tv = max(float(terminal_value_eur_per_kwh), 0.0)
        salvage = np.maximum(levels - soc_min, 0.0) * tv * eff_d
        v_next = salvage.astype(float)
    else:
        v_next = np.zeros(N, dtype=float)

    best_next = np.full((T, N), -1, dtype=np.int16)

    for t in range(T - 1, -1, -1):
        price = float(prices_eur_per_kwh[t])
        v_cur = np.full(N, neg_inf, dtype=float)

        for i in range(N):
            soc_i = levels[i]
            low_soc = max(soc_min, soc_i - max_soc_decrease)
            high_soc = min(soc_max, soc_i + max_soc_increase)

            j_lo = int(np.searchsorted(levels, low_soc, side="left"))
            j_hi = int(np.searchsorted(levels, high_soc, side="right")) - 1
            if j_hi < j_lo:
                continue

            best_val = neg_inf
            best_j = -1

            for j in range(j_lo, j_hi + 1):
                soc_j = levels[j]
                delta_soc = soc_j - soc_i

                if delta_soc > 1e-12:
                    grid_charge = delta_soc / eff_c if eff_c > 0 else 0.0
                    if grid_charge - p_limit_grid_kwh > 1e-9:
                        continue
                    revenue = 0.0
                    cost = grid_charge * price
                    throughput = abs(delta_soc)
                elif delta_soc < -1e-12:
                    soc_out = -delta_soc
                    grid_discharge = soc_out * eff_d
                    if grid_discharge - p_limit_grid_kwh > 1e-9:
                        continue
                    revenue = grid_discharge * price
                    cost = 0.0
                    throughput = abs(delta_soc)
                else:
                    revenue = cost = throughput = 0.0

                immediate = revenue - cost - deg_cost_per_kwh_throughput * throughput
                val = immediate + float(v_next[j])

                if val > best_val:
                    best_val = val
                    best_j = j

            v_cur[i] = best_val
            best_next[t, i] = best_j

        v_next = v_cur

    soc = np.zeros(T + 1, dtype=float)
    grid_ch = np.zeros(T, dtype=float)
    grid_dis = np.zeros(T, dtype=float)

    idx_soc = 0
    soc[0] = levels[idx_soc]
    profit = 0.0

    for t in range(T):
        j = int(best_next[t, idx_soc])
        if j < 0:
            j = idx_soc

        soc_i = levels[idx_soc]
        soc_j = levels[j]
        delta_soc = soc_j - soc_i
        price = float(prices_eur_per_kwh[t])

        if delta_soc > 1e-12:
            grid_charge = delta_soc / eff_c if eff_c > 0 else 0.0
            grid_ch[t] = grid_charge
            profit += -(grid_charge * price) - deg_cost_per_kwh_throughput * delta_soc
        elif delta_soc < -1e-12:
            soc_out = -delta_soc
            grid_discharge = soc_out * eff_d
            grid_dis[t] = grid_discharge
            profit += (grid_discharge * price) - deg_cost_per_kwh_throughput * soc_out

        soc[t + 1] = soc_j
        idx_soc = j

    return float(profit), soc, grid_ch, grid_dis


def optimize_arbitrage_continuous(
    prices_hourly_eur_per_kwh: pd.Series,
    cap_kwh: float,
    p_kw: float,
    dod: float,
    roundtrip_eff: float,
    deg_cost_per_kwh_throughput: float,
    soc_grid_points: int,
    terminal_mode: str,
    terminal_value_eur_per_kwh: float,
) -> dict:
    ser = _to_hourly_prices(prices_hourly_eur_per_kwh).dropna()
    if ser.empty:
        return {"profit_year": 0.0, "throughput_year": 0.0, "cycles_year": 0.0, "schedule_df": pd.DataFrame(), "days_used": 0}

    prices = ser.values.astype(float)
    profit, soc, grid_ch, grid_dis = _dp_optimize_horizon(
        prices_eur_per_kwh=prices,
        cap_kwh=cap_kwh,
        p_kw=p_kw,
        dod=dod,
        roundtrip_eff=roundtrip_eff,
        deg_cost_per_kwh_throughput=deg_cost_per_kwh_throughput,
        soc_grid_points=soc_grid_points,
        terminal_mode=terminal_mode,
        terminal_value_eur_per_kwh=terminal_value_eur_per_kwh,
    )

    deltas = np.diff(soc)
    throughput = float(np.sum(np.abs(deltas)))

    hours = len(ser)
    days = max(hours / 24.0, 1e-9)

    profit_year = (profit / days) * 365.0
    throughput_year = (throughput / days) * 365.0

    usable = max(cap_kwh * float(np.clip(dod, 0.10, 1.0)), 1e-9)
    cycles_year = throughput_year / (2.0 * usable)

    schedule_df = pd.DataFrame({
        "timestamp": ser.index,
        "price_eur_per_kwh": prices,
        "soc_kwh": soc[:-1],
        "soc_next_kwh": soc[1:],
        "grid_charge_kwh": grid_ch,
        "grid_discharge_kwh": grid_dis,
    }).set_index("timestamp")

    return {
        "profit_year": float(profit_year),
        "throughput_year": float(throughput_year),
        "cycles_year": float(cycles_year),
        "schedule_df": schedule_df,
        "days_used": int(round(days)),
    }


def optimize_arbitrage_daily_reset(
    prices_hourly_eur_per_kwh: pd.Series,
    cap_kwh: float,
    p_kw: float,
    dod: float,
    roundtrip_eff: float,
    deg_cost_per_kwh_throughput: float,
    soc_grid_points: int,
) -> dict:
    ser = _to_hourly_prices(prices_hourly_eur_per_kwh).dropna()
    if ser.empty:
        return {"profit_year": 0.0, "throughput_year": 0.0, "cycles_year": 0.0, "schedule_df": pd.DataFrame(), "days_used": 0}

    chunks = []
    total_profit = 0.0
    total_throughput = 0.0

    by_date = ser.groupby(ser.index.date)
    for _, sday in by_date:
        if len(sday) < 2:
            continue
        prices = sday.values.astype(float)
        profit, soc, grid_ch, grid_dis = _dp_optimize_horizon(
            prices_eur_per_kwh=prices,
            cap_kwh=cap_kwh,
            p_kw=p_kw,
            dod=dod,
            roundtrip_eff=roundtrip_eff,
            deg_cost_per_kwh_throughput=deg_cost_per_kwh_throughput,
            soc_grid_points=soc_grid_points,
            terminal_mode="End-SOC = SOC_min",
            terminal_value_eur_per_kwh=0.0,
        )
        deltas = np.diff(soc)
        throughput = float(np.sum(np.abs(deltas)))
        total_profit += profit
        total_throughput += throughput

        df = pd.DataFrame({
            "timestamp": sday.index,
            "price_eur_per_kwh": prices,
            "soc_kwh": soc[:-1],
            "soc_next_kwh": soc[1:],
            "grid_charge_kwh": grid_ch,
            "grid_discharge_kwh": grid_dis,
        }).set_index("timestamp")
        chunks.append(df)

    schedule_df = pd.concat(chunks).sort_index() if chunks else pd.DataFrame()

    days = max(len(by_date), 1)
    profit_year = (total_profit / days) * 365.0
    throughput_year = (total_throughput / days) * 365.0

    usable = max(cap_kwh * float(np.clip(dod, 0.10, 1.0)), 1e-9)
    cycles_year = throughput_year / (2.0 * usable)

    return {
        "profit_year": float(profit_year),
        "throughput_year": float(throughput_year),
        "cycles_year": float(cycles_year),
        "schedule_df": schedule_df,
        "days_used": int(days),
    }


# ================================================================
# 8) KI-SPEICHEREMPFEHLUNG
# ================================================================
def recommend_storage(d: dict, hak_kw: float, objective: str) -> dict:
    pvkwp = float(d["pv"].get("total_kwp", 0.0))
    pvac = float(d["pv"].get("total_ac_kw", 0.0))
    load_building = float(d["wirtschaft"].get("lastgang_jahr", 0.0))
    ev_kwh = float(d["mobilität"].get("total_ev_kwh", 0.0))
    total_load = max(load_building + ev_kwh, 0.0)
    day_load = total_load / 365.0 if total_load > 0 else 0.0

    wb_sum = 0.0
    for lp in d["mobilität"].get("ladepunkte", []):
        wb_sum += float(lp.get("leistung_kw_pro_punkt", 0.0) or 0.0) * float(lp.get("anzahl_punkte", 1) or 1)

    base_kwh = pvkwp * 1.0 if pvkwp > 0 else max(day_load * 0.6, 10.0)
    if objective == "Depotcharging/Lastmanagement":
        base_kwh *= 1.15
    elif objective == "Arbitrage (Stromhandel)":
        base_kwh *= 1.30
    elif objective == "Kombi (Eigenverbrauch + Arbitrage)":
        base_kwh *= 1.20

    min_kwh = day_load * 0.4
    max_kwh = day_load * 2.0 if day_load > 0 else base_kwh * 2.0
    rec_kwh = max(base_kwh, min_kwh)
    rec_kwh = min(rec_kwh, max_kwh) if max_kwh > 0 else rec_kwh

    if objective == "Eigenverbrauch/Autarkie":
        target_c = 0.35 if wb_sum <= 11 else 0.45
    elif objective == "Depotcharging/Lastmanagement":
        target_c = 0.55
    elif objective == "Arbitrage (Stromhandel)":
        target_c = 0.75
    else:
        target_c = 0.55

    rec_kw = rec_kwh * target_c

    reserve = 0.85
    headroom = max(hak_kw * reserve - (pvac + wb_sum), 0.0)
    reasons = []
    if headroom > 0:
        if rec_kw > headroom:
            reasons.append(f"Leistung an HAK-Headroom angepasst (Headroom ~ {headroom:.1f} kW).")
        rec_kw = min(rec_kw, headroom)
    else:
        reasons.append("Kein Headroom (PV-AC + Ladeleistung >= HAK) → Leistung konservativ (EMS/LM Pflicht).")
        rec_kw = min(rec_kw, max(hak_kw * 0.10, 5.0))

    c_rate = (rec_kw / rec_kwh) if rec_kwh > 0 else 0.0
    hv_lv = "Hochvolt" if rec_kwh >= 15 else "Niedervolt"

    ems = False
    if (pvac + wb_sum) > hak_kw:
        ems = True
        reasons.append("PV-AC + Ladeleistung über HAK → Lastmanagement/EMS erforderlich.")
    if wb_sum >= 22:
        ems = True
        reasons.append("Mehrere/leistungsstarke Ladepunkte → dynamisches Lastmanagement sinnvoll.")

    return {
        "rec_kwh": round(float(rec_kwh), 1),
        "rec_kw": round(float(rec_kw), 1),
        "c_rate": round(float(c_rate), 2),
        "typ": hv_lv,
        "ems": ems,
        "reasons": reasons,
    }


# ================================================================
# 9) REPORT / PACKAGE EXPORT (+ index.html im ZIP)
# ================================================================
def collect_attachments(d: dict) -> list[dict]:
    items = []
    w = d.get("wirtschaft", {})
    lg = str(w.get("lastgang_file", "") or "")
    if lg:
        items.append({"type": "Input", "name": "Lastgang CSV", "path": lg})

    arb = d.get("speicher", {}).get("arbitrage", {})
    if arb.get("price_csv_file"):
        items.append({"type": "Arbitrage", "name": "Preis-CSV", "path": arb["price_csv_file"]})
    if arb.get("schedule_file"):
        items.append({"type": "Arbitrage", "name": "Schedule CSV", "path": arb["schedule_file"]})

    for i, f in enumerate(d.get("pv", {}).get("felder", []), start=1):
        p = str(f.get("datasheet", "") or "")
        if p:
            items.append({"type": "PV Modul", "name": f"Feld {i}: {f.get('typ','')}", "path": p})

    for i, wri in enumerate(d.get("pv", {}).get("wr", []), start=1):
        p = str(wri.get("datasheet", "") or "")
        if p:
            items.append({"type": "Wechselrichter", "name": f"WR {i}: {wri.get('modell','')}", "path": p})

    sp = str(d.get("speicher", {}).get("datasheet", "") or "")
    if sp:
        items.append({"type": "Speicher", "name": d.get("speicher", {}).get("hersteller",""), "path": sp})

    for i, lp in enumerate(d.get("mobilität", {}).get("ladepunkte", []), start=1):
        p = str(lp.get("datasheet", "") or "")
        if p:
            items.append({"type": "Ladestation", "name": f"LP {i}: {lp.get('name','')}", "path": p})

    for i, fz in enumerate(d.get("mobilität", {}).get("fuhrpark", []), start=1):
        p = str(fz.get("datasheet", "") or "")
        if p:
            items.append({"type": "Fahrzeug", "name": f"FZ {i}: {fz.get('klasse','')}", "path": p})

    return items


def export_word(report_path: Path, d: dict, sim_kpis: dict):
    if Document is None:
        raise RuntimeError("python-docx nicht installiert (pip install python-docx).")

    doc = Document()
    doc.add_heading("Projektbericht – Profi-Planung OS", level=1)
    doc.add_paragraph(f"Kunde: {d['kunde'].get('name','—')}")
    doc.add_paragraph(f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Gebäudeanschluss", level=2)
    mode = d["tech"].get("hak_mode", "Ampere")
    if mode == "kW":
        doc.add_paragraph(f"HAK: {float(d['tech'].get('hak_kw', 0.0)):.1f} kW (eingetragen) ~ {_hak_ampere_from_kw(float(d['tech'].get('hak_kw',0.0)))} A")
    else:
        doc.add_paragraph(f"HAK: {int(d['tech'].get('hak_ampere',63))} A ~ {_hak_kw_from_ampere(float(d['tech'].get('hak_ampere',63))):.1f} kW")

    doc.add_heading("Dachflächen", level=2)
    for r in d["daecher"]:
        doc.add_paragraph(
            f"- {r.get('name')} | Fläche {float(r.get('flaeche',0)):.1f} m² | Neigung {r.get('neigung')}° | Azimut {r.get('azimut')}° "
            f"| PV manuell: {bool(r.get('pv_kwp_manual_enabled'))} ({float(r.get('pv_kwp_manual',0)):.2f} kWp)"
        )

    doc.add_heading("PV", level=2)
    doc.add_paragraph(
        f"Gesamt kWp: {float(d['pv'].get('total_kwp',0)):.2f} | WR-AC: {float(d['pv'].get('total_ac_kw',0)):.1f} kW | Konzept: {d['pv'].get('konzept','')}"
    )

    y = d.get("pv", {}).get("yield", {}) or {}
    doc.add_paragraph(f"PV-Ertragsmodell: {y.get('mode','Heuristik')} | Standort: {y.get('lat','')} , {y.get('lon','')}")

    doc.add_heading("Ladeinfrastruktur", level=2)
    ac_points = 0
    dc_points = 0
    ac_kw = 0.0
    dc_kw = 0.0
    for lp in d.get("mobilität", {}).get("ladepunkte", []):
        n = int(lp.get("anzahl_punkte", 1) or 1)
        p = float(lp.get("leistung_kw_pro_punkt", 0.0) or 0.0)
        if lp.get("typ") == "DC":
            dc_points += n
            dc_kw += p * n
        else:
            ac_points += n
            ac_kw += p * n
    doc.add_paragraph(f"AC: {ac_points} Punkte / {ac_kw:.1f} kW | DC: {dc_points} Punkte / {dc_kw:.1f} kW")

    doc.add_heading("Speicher", level=2)
    s = d["speicher"]
    doc.add_paragraph(f"{s.get('hersteller','')} | {float(s.get('kap',0)):.1f} kWh | {float(s.get('p',0)):.1f} kW | {s.get('spannung','')}")
    doc.add_paragraph(
        f"Mindestladung: {float(s.get('min_soc_pct',10)):.1f}% | Zyklen (Life): {int(s.get('cycle_life',6000))} | Lebensdauer: {int(s.get('calendar_life_years',15))} Jahre"
    )
    arb = s.get("arbitrage", {})
    doc.add_paragraph(
        f"Arbitrage: {arb.get('enabled', False)} | Mode: {arb.get('mode','Manuell')} | "
        f"Terminal: {arb.get('terminal_mode','End-SOC = SOC_min')} | "
        f"Annual Profit Est: {float(arb.get('annual_profit_est',0)):,.0f} €/a | Cycles Est: {float(arb.get('annual_cycles_est',0)):.1f}/a"
    )

    doc.add_heading("Simulation KPIs", level=2)
    if not sim_kpis:
        doc.add_paragraph("Keine Simulation-KPIs vorhanden.")
    else:
        for k, v in sim_kpis.items():
            doc.add_paragraph(f"- {k}: {v}")

    doc.add_heading("Anhänge / Datenblätter", level=2)
    atts = collect_attachments(d)
    if not atts:
        doc.add_paragraph("Keine Anhänge hinterlegt.")
    else:
        for a in atts:
            doc.add_paragraph(f"- [{a['type']}] {a['name']}: {a['path']}")

    doc.save(str(report_path))


def export_pdf(report_path: Path, d: dict, sim_kpis: dict):
    if canvas is None or A4 is None:
        raise RuntimeError("reportlab nicht installiert (pip install reportlab).")

    c = canvas.Canvas(str(report_path), pagesize=A4)
    w, h = A4

    y = h - 60
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Projektbericht – Profi-Planung OS")
    y -= 30

    c.setFont("Helvetica", 10)
    c.drawString(50, y, f"Kunde: {d['kunde'].get('name','—')}")
    y -= 15
    c.drawString(50, y, f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    y -= 20

    mode = d["tech"].get("hak_mode", "Ampere")
    if mode == "kW":
        line = f"HAK: {float(d['tech'].get('hak_kw', 0.0)):.1f} kW (eingetragen) ~ {_hak_ampere_from_kw(float(d['tech'].get('hak_kw',0.0)))} A"
    else:
        line = f"HAK: {int(d['tech'].get('hak_ampere',63))} A ~ {_hak_kw_from_ampere(float(d['tech'].get('hak_ampere',63))):.1f} kW"
    c.drawString(50, y, line[:120])
    y -= 25

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Dachflächen")
    y -= 18
    c.setFont("Helvetica", 10)
    for r in d["daecher"]:
        line = f"- {r.get('name')} | Fläche {float(r.get('flaeche',0)):.1f} m² | Neigung {r.get('neigung')}° | Az {r.get('azimut')}°"
        c.drawString(55, y, line[:110])
        y -= 14
        if y < 80:
            c.showPage()
            y = h - 60

    y -= 8
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "PV / Speicher / Ladeinfrastruktur")
    y -= 18
    c.setFont("Helvetica", 10)
    c.drawString(55, y, f"PV gesamt: {float(d['pv'].get('total_kwp',0)):.2f} kWp | WR-AC: {float(d['pv'].get('total_ac_kw',0)):.1f} kW | Konzept: {d['pv'].get('konzept','')}"[:120])
    y -= 14

    ycfg = d.get("pv", {}).get("yield", {}) or {}
    c.drawString(55, y, f"PV-Ertragsmodell: {ycfg.get('mode','Heuristik')} | lat/lon: {ycfg.get('lat','')},{ycfg.get('lon','')}"[:120])
    y -= 14

    ac_points = dc_points = 0
    ac_kw = dc_kw = 0.0
    for lp in d.get("mobilität", {}).get("ladepunkte", []):
        n = int(lp.get("anzahl_punkte", 1) or 1)
        p = float(lp.get("leistung_kw_pro_punkt", 0.0) or 0.0)
        if lp.get("typ") == "DC":
            dc_points += n
            dc_kw += p * n
        else:
            ac_points += n
            ac_kw += p * n
    c.drawString(55, y, f"Ladeinfra: AC {ac_points}P/{ac_kw:.1f}kW | DC {dc_points}P/{dc_kw:.1f}kW"[:120])
    y -= 14

    s = d["speicher"]
    arb = s.get("arbitrage", {})
    c.drawString(55, y, f"Speicher: {float(s.get('kap',0)):.1f} kWh / {float(s.get('p',0)):.1f} kW | Arbitrage: {arb.get('enabled',False)} | Terminal: {arb.get('terminal_mode','')}"[:120])
    y -= 18

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Simulation KPIs")
    y -= 18
    c.setFont("Helvetica", 10)
    if not sim_kpis:
        c.drawString(55, y, "Keine Simulation-KPIs vorhanden.")
        y -= 14
    else:
        for k, v in sim_kpis.items():
            c.drawString(55, y, f"- {k}: {v}"[:120])
            y -= 14
            if y < 80:
                c.showPage()
                y = h - 60
                c.setFont("Helvetica", 10)

    y -= 8
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Anhänge / Datenblätter")
    y -= 18
    c.setFont("Helvetica", 9)
    atts = collect_attachments(d)
    if not atts:
        c.drawString(55, y, "Keine Anhänge hinterlegt.")
        y -= 12
    else:
        for a in atts:
            line = f"- [{a['type']}] {a['name']}: {a['path']}"
            c.drawString(55, y, line[:125])
            y -= 12
            if y < 80:
                c.showPage()
                y = h - 60
                c.setFont("Helvetica", 9)

    c.showPage()
    c.save()


def generate_index_html(project_slug: str, d: dict, created_ts: str, docs_root: Path) -> str:
    base = f"projects/{project_slug}/documents"
    files = sorted([p for p in docs_root.rglob("*.*") if p.is_file()], key=lambda x: x.as_posix())
    reports = [p for p in files if "/reports/" in p.as_posix().replace("\\", "/")]
    atts = collect_attachments(d)

    def li(href: str, label: str) -> str:
        return f'<li><a href="{href}">{label}</a></li>'

    html = []
    html.append("<!doctype html><html><head><meta charset='utf-8'>")
    html.append("<title>Projektpaket – Übersicht</title>")
    html.append("<style>body{font-family:Arial,Helvetica,sans-serif;margin:24px;} h1{margin:0 0 8px;} .small{color:#666;font-size:12px;} code{background:#f4f4f4;padding:2px 4px;border-radius:4px;}</style>")
    html.append("</head><body>")
    html.append(f"<h1>Projektpaket – {project_slug}</h1>")
    html.append(f"<div class='small'>Erstellt: {created_ts} | Kunde: <b>{d.get('kunde',{}).get('name','—')}</b></div>")

    mode = d.get("tech", {}).get("hak_mode", "Ampere")
    if mode == "kW":
        hak_line = f"{float(d['tech'].get('hak_kw',0.0)):.1f} kW (~ {_hak_ampere_from_kw(float(d['tech'].get('hak_kw',0.0)))} A)"
    else:
        hak_line = f"{int(d['tech'].get('hak_ampere',63))} A (~ {_hak_kw_from_ampere(float(d['tech'].get('hak_ampere',63))):.1f} kW)"
    html.append(f"<p><b>Gebäudeanschluss:</b> {hak_line}</p>")

    html.append("<h2>Reports</h2><ul>")
    if not reports:
        html.append("<li>Keine Reports gefunden.</li>")
    else:
        for p in reports:
            rel = p.as_posix().replace(docs_root.as_posix().replace("\\", "/"), base).replace("\\", "/")
            html.append(li(rel, p.name))
    html.append("</ul>")

    html.append("<h2>Anlagen / Datenblätter</h2><ul>")
    if not atts:
        html.append("<li>Keine Anhänge hinterlegt.</li>")
    else:
        for a in atts:
            p = a.get("path", "")
            rel = p.replace("\\", "/")
            html.append(li(rel, f"[{a.get('type','')}] {a.get('name','')}"))
    html.append("</ul>")

    html.append("<h2>Alle Dateien im Paket</h2><ul>")
    for p in files:
        rel = p.as_posix().replace(docs_root.as_posix().replace("\\", "/"), base).replace("\\", "/")
        html.append(li(rel, rel))
    html.append("</ul>")

    html.append("<p class='small'>Hinweis: Öffne index.html nach dem Entpacken im Browser.</p>")
    html.append("</body></html>")
    return "\n".join(html)


def build_zip_package(zip_path: Path, files: list[Path], extra_bytes: dict[str, bytes] | None = None) -> Path:
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in files:
            try:
                if f.exists() and f.is_file():
                    z.write(f, arcname=str(f).replace("\\", "/"))
            except Exception:
                pass
        if extra_bytes:
            for arcname, b in extra_bytes.items():
                try:
                    z.writestr(arcname, b)
                except Exception:
                    pass
    return zip_path


def gather_project_package_files(docs_root: Path, state_file: Path) -> list[Path]:
    files: list[Path] = []
    if state_file.exists():
        files.append(state_file)

    for folder in [
        docs_root / "reports",
        docs_root / "datasheets",
        docs_root / "inputs",
        docs_root / "arbitrage",
    ]:
        if folder.exists():
            for p in folder.rglob("*.*"):
                if p.is_file():
                    files.append(p)

    seen, uniq = set(), []
    for p in files:
        sp = str(p)
        if sp not in seen:
            seen.add(sp)
            uniq.append(p)
    return uniq


# ================================================================
# 10) APP START / LOAD PROJECT
# ================================================================
if "active_slug" not in st.session_state:
    st.error("⚠️ Kein aktives Projekt. Bitte über das Dashboard starten.")
    st.stop()

P_SLUG = st.session_state["active_slug"]
BASE, DOCS, DATASHEETS, INPUTS, REPORTS, ARB_DIR = project_paths(P_SLUG)

if "project_data" not in st.session_state or st.session_state.get("_loaded_slug") != P_SLUG:
    d, STATE_FILE = load_project(P_SLUG)
    st.session_state["project_data"] = d
    st.session_state["_loaded_slug"] = P_SLUG
    st.session_state["_last_saved_hash"] = dict_hash(d)
else:
    d = st.session_state["project_data"]
    STATE_FILE = Path("projects") / P_SLUG / "state.json"
    deep_merge(d, DEFAULTS)
    migrate_legacy(d)

hak_kw = hak_p_max_kw_from_state(d)


# ================================================================
# SIDEBAR
# ================================================================
with st.sidebar:
    st.header("🎯 Projekt-Sektoren")
    current_scope = [x for x in d.get("scope", []) if x in SCOPE_OPTIONS] or DEFAULTS["scope"]
    d["scope"] = st.multiselect("Themen zur Planung:", SCOPE_OPTIONS, default=current_scope)

    st.divider()
    st.header("💼 Lokale Sicherung (ZIP)")
    st.caption("Browser-Apps dürfen nicht direkt auf deinen PC schreiben. Lösung: ZIP herunterladen (lokal speichern) & später wieder hochladen.")
    zbytes = export_project_zip_bytes(P_SLUG)
    st.download_button("⬇️ Projekt-ZIP herunterladen", data=zbytes, file_name=f"{P_SLUG}.zip", use_container_width=True)

    up_zip = st.file_uploader("ZIP importieren", type=["zip"], key="zip_import_sidebar")
    if up_zip is not None:
        try:
            new_slug = import_project_zip_bytes(up_zip.getvalue())
            if new_slug:
                st.success(f"Import ok: {new_slug}")
                st.session_state["active_slug"] = new_slug
                st.session_state["_loaded_slug"] = None
                st.rerun()
            else:
                st.warning("Import durchgeführt, aber kein Projektordner erkannt.")
        except Exception as e:
            st.error(str(e))

    st.divider()
    st.header("⚡ Gebäudeanschluss")
    d["tech"]["hak_mode"] = st.radio(
        "Eingabeart",
        HAK_MODE_OPTIONS,
        index=safe_index(HAK_MODE_OPTIONS, d["tech"].get("hak_mode", "Ampere"), 0),
        horizontal=True
    )

    if d["tech"]["hak_mode"] == "Ampere":
        d["tech"]["hak_ampere"] = st.selectbox(
            "HAK Größe (Ampere)",
            HAK_OPTIONS,
            index=safe_index(HAK_OPTIONS, int(d["tech"].get("hak_ampere", 63)), 2),
        )
        hak_kw = _hak_kw_from_ampere(float(d["tech"]["hak_ampere"]))
        d["tech"]["hak_kw"] = round(float(hak_kw), 1)
    else:
        d["tech"]["hak_kw"] = st.number_input(
            "HAK Größe (kW)",
            1.0, 20000.0,
            float(d["tech"].get("hak_kw", _hak_kw_from_ampere(63))),
            step=1.0
        )
        hak_kw = float(d["tech"]["hak_kw"])
        d["tech"]["hak_ampere"] = _hak_ampere_from_kw(hak_kw)

    st.metric("Max. Belastbarkeit (theoret.)", f"{hak_kw:.1f} kW")
    st.caption(f"≈ {int(d['tech'].get('hak_ampere',63))} A bei 400V/3~")

    st.divider()
    if st.button("💾 Projektstand speichern", use_container_width=True):
        save_project(STATE_FILE, d)
        st.session_state["_last_saved_hash"] = dict_hash(d)
        st.toast("Daten lokal gesichert!")


# ================================================================
# TITEL + TABS
# ================================================================
st.title(f"Fachplanung: {d['kunde'].get('name','—')}")

tab_names = []
if "PV-System" in d["scope"]:
    tab_names.append("☀️ PV & Dächer")
if "Speicher" in d["scope"]:
    tab_names.append("🔋 Speicher & Arbitrage")
if "Ladeinfrastruktur" in d["scope"]:
    tab_names.append("🔌 E-Mobilität")
tab_names.append("📈 Simulation & ROI")
tab_names.append("🧾 Report & Dateien")

tabs = st.tabs(tab_names)
ti = 0

# ================================================================
# TAB: PV & DÄCHER  (unverändert – gekürzt)
# ================================================================
if "PV-System" in d["scope"]:
    with tabs[ti]:
        st.header("☀️ PV & Dachflächen")
        # (Dieser Bereich bleibt identisch zu deiner Version; hier aus Platzgründen nicht erneut kommentiert.)
        # >>> START PV & DÄCHER (dein Original) >>>
        with st.container(border=True):
            st.subheader("Dachflächen (inkl. PV-kWp manuell je Dach)")
            c_add, c_info = st.columns([1, 3])
            with c_add:
                if st.button("➕ Dachfläche hinzufügen"):
                    rid = f"roof-{len(d['daecher'])+1}"
                    d["daecher"].append({
                        "id": rid,
                        "name": f"Dachfläche {len(d['daecher'])+1}",
                        "form": "Satteldach",
                        "breite": 10.0,
                        "tiefe": 6.0,
                        "flaeche": 60.0,
                        "flaeche_auto": True,
                        "neigung": 35,
                        "azimut": 0,
                        "nutzfaktor": 0.80,
                        "module_area_m2": 2.0,
                        "pv_kwp_manual_enabled": False,
                        "pv_kwp_manual": 0.0,
                        "hinweis": "",
                    })
                    st.rerun()
            with c_info:
                st.caption("Wenn du pro Dach keine Modulliste pflegen willst: „PV-kWp manuell“ aktivieren.")

            for i, r in enumerate(list(d["daecher"])):
                with st.expander(f"🏠 {r.get('name','Dachfläche')} (#{i+1})", expanded=(i == 0)):
                    c1, c2, c3, c4 = st.columns([1.4, 1, 1, 1])
                    r["name"] = c1.text_input("Name", r.get("name", f"Dachfläche {i+1}"), key=f"roof_name_{i}")
                    r["form"] = c2.selectbox("Dachform", DACH_FORMS, index=safe_index(DACH_FORMS, r.get("form", "Satteldach"), 0), key=f"roof_form_{i}")
                    r["neigung"] = c3.number_input("Neigung (°)", 0, 90, int(r.get("neigung", 35)), key=f"roof_tilt_{i}")
                    r["azimut"] = c4.number_input("Azimut (°)", -180, 180, int(r.get("azimut", 0)), key=f"roof_az_{i}")

                    c5, c6, c7, c8 = st.columns([1, 1, 1, 1])
                    r["breite"] = c5.number_input("Breite (m)", 0.0, 5000.0, float(r.get("breite", 10.0)), key=f"roof_w_{i}")
                    r["tiefe"] = c6.number_input("Tiefe/Höhe (m)", 0.0, 5000.0, float(r.get("tiefe", 6.0)), key=f"roof_d_{i}")
                    r["flaeche_auto"] = c7.checkbox("Fläche auto", value=bool(r.get("flaeche_auto", True)), key=f"roof_auto_{i}")

                    auto_area = float(roof_area(r))
                    max_area_ui = max(10_000.0, min(auto_area + 1.0, PLAUS_ROOF_AREA_HARD))

                    if r["flaeche_auto"]:
                        r["flaeche"] = auto_area
                        c8.number_input("Fläche (m²)", min_value=0.0, max_value=max_area_ui, value=float(r["flaeche"]), disabled=True, key=f"roof_area_{i}")
                    else:
                        r["flaeche"] = c8.number_input("Fläche (m²)", min_value=0.0, max_value=PLAUS_ROOF_AREA_HARD, value=float(r.get("flaeche", auto_area)), key=f"roof_area_{i}")

                    wv = float(r.get("breite", 0.0) or 0.0)
                    dv = float(r.get("tiefe", 0.0) or 0.0)
                    av = float(r.get("flaeche", 0.0) or 0.0)
                    if wv > PLAUS_ROOF_W_MAX or dv > PLAUS_ROOF_D_MAX:
                        st.warning(f"⚠️ Plausibilität: sehr große Dachabmessungen (Breite {wv:.1f} m / Tiefe {dv:.1f} m).")
                    if av > PLAUS_ROOF_AREA_WARN:
                        st.warning(f"⚠️ Plausibilität: Dachfläche {av:,.0f} m² ist sehr groß.")
                    if av >= PLAUS_ROOF_AREA_HARD:
                        st.error("🚨 Dachfläche erreicht das UI-Limit. Bitte Werte prüfen/aufteilen in mehrere Dächer.")

                    c9, c10, c11 = st.columns([1, 1, 1])
                    r["nutzfaktor"] = c9.slider("Nutzfaktor", 0.30, 0.95, float(r.get("nutzfaktor", 0.80)), 0.01, key=f"roof_nf_{i}")
                    r["module_area_m2"] = c10.number_input("Modulfläche (m²)", 0.5, 5.0, float(r.get("module_area_m2", 2.0)), 0.1, key=f"roof_ma_{i}")
                    r["pv_kwp_manual_enabled"] = c11.checkbox("PV-kWp manuell", value=bool(r.get("pv_kwp_manual_enabled", False)), key=f"roof_pvman_{i}")
                    if r["pv_kwp_manual_enabled"]:
                        r["pv_kwp_manual"] = st.number_input("PV-Leistung (kWp) dieses Dach", 0.0, 500000.0, float(r.get("pv_kwp_manual", 0.0)), key=f"roof_pvman_val_{i}")

                    cap_mod = roof_module_capacity(r)
                    used_mod = modules_from_fields(d["pv"]["felder"], r["id"])
                    pv_dach_kwp = float(r.get("pv_kwp_manual", 0.0)) if bool(r.get("pv_kwp_manual_enabled")) else pv_kwp_from_fields(d["pv"]["felder"], r["id"])

                    kA, kB, kC = st.columns(3)
                    kA.metric("Kapazität grob", f"{cap_mod:,} Module")
                    kB.metric("Zugeordnet", f"{used_mod:,} Module")
                    kC.metric("PV-kWp (Dach)", f"{pv_dach_kwp:.2f} kWp")

                    r["hinweis"] = st.text_input("Hinweis (optional)", r.get("hinweis", ""), key=f"roof_note_{i}")

        st.divider()

        st.subheader("🟦 PV-Modulfelder (mit Dach-Zuordnung + Löschen + Datenblatt)")
        roof_map = {r["id"]: r.get("name", r["id"]) for r in d["daecher"]}
        roof_ids = list(roof_map.keys())

        if st.button("➕ Modulfeld hinzufügen"):
            d["pv"]["felder"].append({
                "roof_id": roof_ids[0] if roof_ids else "roof-1",
                "hersteller": "",
                "typ": "440W Standard",
                "watt_pro_modul": 440,
                "anzahl_module": 12,
                "datasheet": "",
            })

        if d["pv"]["felder"]:
            for i, f in enumerate(list(d["pv"]["felder"])):
                with st.expander(f"Feld {i+1}: {f.get('typ','') or '—'}", expanded=(i == 0)):
                    c1, c2, c3 = st.columns([1.2, 1, 1])
                    f["roof_id"] = c1.selectbox("Dach", options=roof_ids, index=safe_index(roof_ids, f.get("roof_id", roof_ids[0] if roof_ids else "roof-1"), 0), key=f"pvf_roof_{i}")
                    f["hersteller"] = c2.text_input("Hersteller", f.get("hersteller", ""), key=f"pvf_h_{i}")
                    f["typ"] = c3.text_input("Typ", f.get("typ", ""), key=f"pvf_t_{i}")

                    c4, c5, c6 = st.columns([1, 1, 1])
                    f["watt_pro_modul"] = c4.number_input("Watt/Modul", 0, 2000, int(f.get("watt_pro_modul", 0) or 0), key=f"pvf_w_{i}")
                    f["anzahl_module"] = c5.number_input("Anzahl Module", 0, 500000, int(f.get("anzahl_module", 0) or 0), key=f"pvf_n_{i}")

                    with c6:
                        up = st.file_uploader("Datenblatt (PDF)", type=None, key=f"pvf_ds_up_{i}")
                        if up is not None:
                            f["datasheet"] = save_uploaded_file(DATASHEETS / "pv_module", up, prefix="pvmod_")
                            st.success("Gespeichert.")
                        if f.get("datasheet"):
                            st.code(f["datasheet"])

                    if st.button("🗑️ Modulfeld löschen", key=f"pvf_del_{i}"):
                        d["pv"]["felder"].pop(i)
                        st.rerun()
        else:
            st.info("Noch keine PV-Felder angelegt.")

        st.divider()
        st.subheader("🚥 Wechselrichter (Löschen + Datenblatt)")
        if st.button("➕ Wechselrichter hinzufügen"):
            d["pv"]["wr"].append({"modell": "", "ac_kw": 10.0, "datasheet": ""})

        if d["pv"]["wr"]:
            for j, wri in enumerate(list(d["pv"]["wr"])):
                with st.container(border=True):
                    c1, c2, c3 = st.columns([1.5, 1, 1])
                    wri["modell"] = c1.text_input("Modell", wri.get("modell", ""), key=f"wr_m_{j}")
                    wri["ac_kw"] = c2.number_input("AC-Nennleistung (kW)", 0.0, 5_000_000.0, float(wri.get("ac_kw", 0.0) or 0.0), key=f"wr_p_{j}")
                    with c3:
                        up = st.file_uploader("Datenblatt", type=None, key=f"wr_ds_up_{j}")
                        if up is not None:
                            wri["datasheet"] = save_uploaded_file(DATASHEETS / "wechselrichter", up, prefix="wr_")
                            st.success("Gespeichert.")
                        if wri.get("datasheet"):
                            st.code(wri["datasheet"])
                    if st.button("🗑️ Wechselrichter löschen", key=f"wr_del_{j}"):
                        d["pv"]["wr"].pop(j)
                        st.rerun()
        else:
            st.info("Noch keine Wechselrichter angelegt.")

        st.divider()
        d["pv"]["konzept"] = st.radio(
            "Einspeisekonzept",
            PV_KONZEPTE,
            index=safe_index(PV_KONZEPTE, d["pv"].get("konzept", "Überschusseinspeisung"), 0),
            horizontal=True
        )
        if d["pv"]["konzept"] == "Volleinspeisung":
            st.warning("⚠️ Hinweis: Volleinspeisung benötigt häufig separaten Zählerplatz / Messkonzept.")

        d["pv"]["total_kwp"] = pv_kwp_total(d)
        d["pv"]["total_ac_kw"] = wr_ac_total(d)
        c1, c2 = st.columns(2)
        c1.metric("PV gesamt", f"{d['pv']['total_kwp']:.2f} kWp")
        c2.metric("WR-AC gesamt", f"{d['pv']['total_ac_kw']:.1f} kW")

    ti += 1

# ================================================================
# TAB: SPEICHER & ARBITRAGE (unverändert – bereits vollständig oben)
# ================================================================
if "Speicher" in d["scope"]:
    # Für Übersichtlichkeit: du kannst hier einfach deinen vorhandenen Speicher-Tab aus deiner Version lassen.
    # In diesem Export-Script ist er aus Platzgründen NICHT erneut eingefügt.
    with tabs[ti]:
        st.info("🔋 Speicher-Tab: Bitte ersetze diesen Platzhalter durch deinen bestehenden Speicher/Arbitrage-Tab (aus deiner Version).")
    ti += 1

# ================================================================
# TAB: E-MOBILITÄT (unverändert – bitte deinen bestehenden Tab einsetzen)
# ================================================================
if "Ladeinfrastruktur" in d["scope"]:
    with tabs[ti]:
        st.info("🔌 E‑Mobilität‑Tab: Bitte ersetze diesen Platzhalter durch deinen bestehenden E‑Mobilität‑Tab (aus deiner Version).")
    ti += 1

# ================================================================
# TAB: SIMULATION & ROI (hier: PVGIS optional)
# ================================================================
with tabs[ti]:
    st.header("📈 Simulation & ROI")

    with st.expander("Preise & Annahmen", expanded=True):
        c1, c2, c3, c4 = st.columns(4)
        d["wirtschaft"]["strompreis"] = c1.number_input("Strompreis Netz (€/kWh)", 0.05, 2.50, float(d["wirtschaft"].get("strompreis", 0.35)))
        d["wirtschaft"]["einspeise_v"] = c2.number_input("Einspeisevergütung (€/kWh)", 0.00, 1.00, float(d["wirtschaft"].get("einspeise_v", 0.08)))
        d["wirtschaft"]["dieselpreis"] = c3.number_input("Dieselpreis (€/L)", 0.50, 5.00, float(d["wirtschaft"].get("dieselpreis", 1.70)))

        # --- PV-Ertrag Modell (neu) ---
        y = d["pv"].setdefault("yield", DEFAULTS["pv"]["yield"])
        y["mode"] = c4.selectbox("PV-Ertragsmodell", PV_YIELD_MODE_OPTIONS, index=safe_index(PV_YIELD_MODE_OPTIONS, y.get("mode", PV_YIELD_MODE_OPTIONS[0]), 0))

        if y["mode"].startswith("Heuristik"):
            region = st.selectbox("Region (PV-Ertrag Modell)", list(REGIONS.keys()), index=safe_index(list(REGIONS.keys()), y.get("region", "Mitte"), 1))
            y["region"] = region
            st.caption("Heuristik: Jahresertrag ≈ kWh/kWp/a (Nord/Mitte/Süd).")
        else:
            st.caption("PVGIS: reale Standortdaten (EU JRC). Azimut: 0=süd, 90=west, -90=ost.")
            loc_q = st.text_input("Standort (Adresse / PLZ / Ort)", value=str(y.get("location_query", "")))
            y["location_query"] = loc_q

            cL1, cL2, cL3 = st.columns([1, 1, 2])
            with cL1:
                if st.button("📍 Koordinaten suchen"):
                    try:
                        hit = geocode_nominatim(loc_q)
                        y["lat"] = float(hit["lat"])
                        y["lon"] = float(hit["lon"])
                        st.success(hit.get("display_name", "Gefunden"))
                    except Exception as e:
                        st.warning(str(e))
            y["lat"] = float(cL2.number_input("Latitude", -90.0, 90.0, float(y.get("lat", 51.1657)), step=0.0001))
            y["lon"] = float(cL3.number_input("Longitude", -180.0, 180.0, float(y.get("lon", 10.4515)), step=0.0001))

            cY1, cY2, cY3, cY4 = st.columns(4)
            y["loss_pct"] = float(cY1.number_input("Systemverluste (%)", 0.0, 40.0, float(y.get("loss_pct", 14.0)), step=0.5))
            y["ref_year"] = int(cY2.number_input("PVGIS Referenzjahr", 2005, 2023, int(y.get("ref_year", 2020)), step=1))
            y["mountingplace"] = cY3.selectbox("Montage", ["free", "building"], index=0 if str(y.get("mountingplace", "free")) == "free" else 1)
            y["use_roof_angles"] = cY4.checkbox("Je Dach Neigung/Azimut nutzen", value=bool(y.get("use_roof_angles", True)))

            if not y["use_roof_angles"]:
                cT1, cT2 = st.columns(2)
                y["tilt_default"] = float(cT1.number_input("Tilt Default (°)", 0.0, 90.0, float(y.get("tilt_default", 30.0)), step=1.0))
                y["azimuth_default"] = float(cT2.number_input("Azimut Default (°)", -180.0, 180.0, float(y.get("azimuth_default", 0.0)), step=1.0))

            with st.expander("PVGIS Test (kWh/kWp/a)"):
                try:
                    test = pvgis_power_per_kwp_hourly(
                        lat=float(y["lat"]), lon=float(y["lon"]),
                        year=int(y["ref_year"]),
                        tilt=float(y.get("tilt_default", 30.0)),
                        azimuth=float(y.get("azimuth_default", 0.0)),
                        loss_pct=float(y.get("loss_pct", 14.0)),
                        mountingplace=str(y.get("mountingplace", "free")),
                        raddatabase=str(y.get("raddatabase", "") or None)
                    )
                    # Stundenleistung in kW -> kWh/a (≈ Summe kW * 1h)
                    kwh_per_kwp = float(test.sum())
                    st.metric("PVGIS Ertrag (grob)", f"{kwh_per_kwp:,.0f} kWh/kWp/a")
                    st.line_chart(test.head(24 * 14))
                except Exception as e:
                    st.warning(str(e))

        c5, c6, c7 = st.columns(3)
        d["wirtschaft"]["capex_pv_eur_per_kwp"] = c5.number_input("CAPEX PV (€/kWp)", 200.0, 4000.0, float(d["wirtschaft"].get("capex_pv_eur_per_kwp", 1100.0)))
        d["wirtschaft"]["capex_bat_eur_per_kwh"] = c6.number_input("CAPEX Speicher (€/kWh)", 50.0, 2000.0, float(d["wirtschaft"].get("capex_bat_eur_per_kwh", 450.0)))
        d["wirtschaft"]["capex_bat_eur_per_kw"] = c7.number_input("CAPEX Speicher (€/kW)", 0.0, 2000.0, float(d["wirtschaft"].get("capex_bat_eur_per_kw", 200.0)))

    pvkwp = float(d["pv"].get("total_kwp", 0.0))
    if pvkwp <= 0:
        st.info("Bitte PV konfigurieren.")
    else:
        s_load = load_lastgang_from_state(d)

        resolution = str(d["wirtschaft"].get("lastgang_resolution", "15min") or "15min")
        if resolution not in SIM_RES_OPTIONS:
            resolution = "15min"

        if s_load is None:
            st.warning("Kein Lastgang in state.json gefunden → nutze synthetisches Lastprofil (Fallback).")
            freq = "15T" if resolution == "15min" else "H"
            idx = pd.date_range("2026-01-01", "2027-01-01", freq=freq, inclusive="left", tz=None)
            seed = stable_seed(P_SLUG, y.get("region", "Mitte"), pvkwp, "fallback")
            s_load = build_synth_load(idx, float(d["wirtschaft"].get("lastgang_jahr", 5000)), seed)
        else:
            idx = s_load.index
            d["wirtschaft"]["lastgang_jahr"] = float(s_load.sum())

        seed = stable_seed(P_SLUG, y.get("region", "Mitte"), pvkwp, int(len(idx)))

        # --- PV-Profil wählen: Heuristik oder PVGIS ---
        pv_model_used = "Heuristik"
        try:
            if str(y.get("mode", "")).startswith("PVGIS"):
                pv_kwh_step = build_pvgis_pv_energy_profile(idx, d)
                pv_model_used = "PVGIS (Standort)"
            else:
                pv_kwh_step = build_synth_pv_profile(idx, pvkwp, str(y.get("region", "Mitte")), seed)
                pv_model_used = f"Heuristik ({y.get('region','Mitte')})"
        except Exception as e:
            pv_kwh_step = build_synth_pv_profile(idx, pvkwp, str(y.get("region", "Mitte")), seed)
            pv_model_used = f"Heuristik (Fallback) – PVGIS Fehler: {e}"

        st.caption(f"PV-Ertragsmodell aktiv: **{pv_model_used}**")

        ev_kwh_year = float(d.get("mobilität", {}).get("total_ev_kwh", 0.0))
        ev_kwh_step = build_ev_profile(idx, ev_kwh_year, seed)

        load_kwh_step = s_load + ev_kwh_step

        s = d["speicher"]
        bat_kwh = float(s.get("kap", 0.0))
        bat_kw = float(s.get("p", 0.0))
        eta = float(s.get("eta_roundtrip", 0.92))
        min_soc_pct = float(s.get("min_soc_pct", 10.0))

        sim = simulate_timeseries(
            pv_kwh_step,
            load_kwh_step,
            bat_kwh,
            bat_kw,
            roundtrip_eff=eta,
            min_soc_pct=min_soc_pct,
        )

        st.subheader("Jahresverlauf (letzte 7 Tage)")
        tail_n = (24 * 7) if resolution == "hour" else (24 * 7 * 4)
        st.line_chart(sim[["PV", "Last", "Served"]].tail(tail_n))

        total_gen = float(sim["PV"].sum())
        total_load = float(sim["Last"].sum())
        total_served = float(sim["Served"].sum())
        total_import = float(sim["Import"].sum())
        total_export = float(sim["Export"].sum())

        autarkie = (total_served / total_load * 100.0) if total_load > 0 else 0.0
        eigenq = (total_served / total_gen * 100.0) if total_gen > 0 else 0.0

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("PV-Ertrag", f"{total_gen:,.0f} kWh/a")
        c2.metric("Autarkie", f"{autarkie:.1f} %")
        c3.metric("Eigenverbrauchsquote", f"{eigenq:.1f} %")
        c4.metric("Netzbezug", f"{total_import:,.0f} kWh/a")
        st.caption(f"Netzeinspeisung: {total_export:,.0f} kWh/a")

        with st.expander("SoC / Import / Export"):
            st.line_chart(sim[["SoC"]].tail(tail_n))
            st.line_chart(sim[["Import", "Export"]].tail(tail_n))

        strompreis = float(d["wirtschaft"]["strompreis"])
        feedin = float(d["wirtschaft"]["einspeise_v"])

        baseline_cost = total_load * strompreis
        system_cost = (total_import * strompreis) - (total_export * feedin)
        annual_savings = baseline_cost - system_cost

        capex_pv = pvkwp * float(d["wirtschaft"]["capex_pv_eur_per_kwp"])
        capex_bat = bat_kwh * float(d["wirtschaft"]["capex_bat_eur_per_kwh"]) + bat_kw * float(d["wirtschaft"]["capex_bat_eur_per_kw"])
        opex = capex_pv * float(d["wirtschaft"]["opex_pv_pct"]) + capex_bat * float(d["wirtschaft"]["opex_bat_pct"])
        net_savings = annual_savings - opex

        st.subheader("💶 ROI (grob)")
        c5, c6, c7, c8 = st.columns(4)
        c5.metric("Baseline", f"{baseline_cost:,.0f} €/a")
        c6.metric("Mit PV/Speicher", f"{system_cost:,.0f} €/a")
        c7.metric("Ersparnis brutto", f"{annual_savings:,.0f} €/a")
        c8.metric("Ersparnis netto (−OPEX)", f"{net_savings:,.0f} €/a")

        capex_total = capex_pv + capex_bat
        if net_savings > 0:
            st.info(f"Grobe Amortisation: **{capex_total / net_savings:.1f} Jahre**")
        else:
            st.warning("Netto-Ersparnis ≤ 0 → Annahmen/Dimensionierung prüfen.")

        st.session_state["_last_sim_kpis"] = {
            "PV-Ertrag (kWh/a)": f"{total_gen:,.0f}",
            "Verbrauch (kWh/a)": f"{total_load:,.0f}",
            "Autarkie (%)": f"{autarkie:.1f}",
            "Eigenverbrauchsquote (%)": f"{eigenq:.1f}",
            "Netzbezug (kWh/a)": f"{total_import:,.0f}",
            "Einspeisung (kWh/a)": f"{total_export:,.0f}",
            "Ersparnis netto (€/a)": f"{net_savings:,.0f}",
        }

ti += 1

# ================================================================
# TAB: REPORT & DATEIEN (unverändert – in deiner Version vorhanden)
# ================================================================
with tabs[ti]:
    st.info("🧾 Report‑Tab: Bitte ersetze diesen Platzhalter durch deinen bestehenden Report/ZIP‑Tab (aus deiner Version).")

# ================================================================
# AUTOSAVE
# ================================================================
save_if_changed(STATE_FILE, d)
